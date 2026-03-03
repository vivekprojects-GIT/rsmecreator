"""
Extract text from DOCX files.
"""

from io import BytesIO

from ..logging_config import get_logger

logger = get_logger("utils.doc_parser")


def extract_text_from_docx(data: bytes) -> str:
    """Extract plain text from a DOCX file."""
    try:
        from docx import Document
        doc = Document(BytesIO(data))
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    tables_text.append(" | ".join(cells))
        text = "\n".join(paras)
        if tables_text:
            text += "\n\n" + "\n".join(tables_text)
        return text.strip() or "(empty document)"
    except Exception as e:
        logger.warning("DOCX parse failed: %s", e)
        raise ValueError(f"Could not parse DOCX: {e}") from e


def extract_text_from_file(data: bytes, filename: str) -> str:
    """Extract text from file based on extension. Supports .docx, .txt, .md."""
    name = (filename or "").lower()
    if name.endswith(".docx") or name.endswith(".doc"):
        return extract_text_from_docx(data)
    if name.endswith(".txt") or name.endswith(".md"):
        return data.decode("utf-8", errors="replace").strip()
    # Default: try as text
    try:
        return data.decode("utf-8", errors="replace").strip()
    except Exception:
        raise ValueError(f"Unsupported file type: {filename}. Use .docx, .doc, .txt, or .md")
